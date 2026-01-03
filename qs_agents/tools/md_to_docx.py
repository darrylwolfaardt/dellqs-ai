"""
Markdown to DOCX Converter

Converts markdown files to Word documents (.docx) format.
Handles common markdown elements: headings, paragraphs, bold, italic,
lists, code blocks, and tables.

Usage:
    python -m qs_agents.tools.md_to_docx input.md [output.docx]

    Or as a module:
        from qs_agents.tools.md_to_docx import convert_md_to_docx
        convert_md_to_docx("input.md", "output.docx")
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def parse_inline_formatting(paragraph, text: str):
    """Parse inline markdown formatting (bold, italic, code) and add to paragraph."""
    # Pattern to match bold, italic, code, or plain text
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))'

    for match in re.finditer(pattern, text):
        if match.group(2):  # Bold italic ***text***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # Bold **text**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):  # Italic *text*
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # Code `text`
            run = paragraph.add_run(match.group(5))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif match.group(6):  # Plain text
            paragraph.add_run(match.group(6))


def convert_md_to_docx(input_path: str, output_path: str = None) -> str:
    """
    Convert a markdown file to a Word document.

    Args:
        input_path: Path to the markdown file
        output_path: Path for the output docx file (optional, defaults to same name with .docx)

    Returns:
        Path to the created docx file
    """
    input_path = Path(input_path)
    if output_path is None:
        output_path = input_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)

    # Read markdown content
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Process line by line
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_para = doc.add_paragraph()
                code_para.style = 'No Spacing'
                for code_line in code_block_content:
                    run = code_para.add_run(code_line + '\n')
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                code_block_content = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Handle tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []

            # Skip separator lines (|---|---|)
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue

            # Parse table row
            cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table, create it
            if table_rows:
                num_cols = max(len(row) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            parse_inline_formatting(cell.paragraphs[0], cell_text)

                doc.add_paragraph()  # Add space after table
            table_rows = []
            in_table = False
            # Don't increment i, process current line

        # Handle headings
        if line.startswith('#'):
            match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                text = match.group(2)
                heading = doc.add_heading(level=min(level, 9))
                parse_inline_formatting(heading, text)
                i += 1
                continue

        # Handle horizontal rules
        if re.match(r'^[-*_]{3,}\s*$', line):
            # Add a paragraph with a bottom border
            para = doc.add_paragraph()
            para.add_run('_' * 50)
            i += 1
            continue

        # Handle unordered lists
        if re.match(r'^[\s]*[-*+]\s+', line):
            match = re.match(r'^([\s]*)[-*+]\s+(.+)$', line)
            if match:
                indent = len(match.group(1))
                text = match.group(2)
                para = doc.add_paragraph(style='List Bullet')
                para.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
                parse_inline_formatting(para, text)
            i += 1
            continue

        # Handle ordered lists
        if re.match(r'^[\s]*\d+\.\s+', line):
            match = re.match(r'^([\s]*)\d+\.\s+(.+)$', line)
            if match:
                indent = len(match.group(1))
                text = match.group(2)
                para = doc.add_paragraph(style='List Number')
                para.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
                parse_inline_formatting(para, text)
            i += 1
            continue

        # Handle blockquotes
        if line.startswith('>'):
            text = line.lstrip('>').strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            para.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else 'Normal'
            parse_inline_formatting(para, text)
            i += 1
            continue

        # Handle regular paragraphs
        stripped = line.strip()
        if stripped:
            para = doc.add_paragraph()
            parse_inline_formatting(para, stripped)

        i += 1

    # Handle any remaining table
    if in_table and table_rows:
        num_cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'

        for row_idx, row_data in enumerate(table_rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    parse_inline_formatting(cell.paragraphs[0], cell_text)

    # Save document
    doc.save(str(output_path))
    return str(output_path)


def convert_multiple(input_paths: list, output_dir: str = None) -> list:
    """
    Convert multiple markdown files to docx.

    Args:
        input_paths: List of markdown file paths
        output_dir: Optional output directory (defaults to same directory as input)

    Returns:
        List of created docx file paths
    """
    results = []
    for input_path in input_paths:
        input_path = Path(input_path)
        if output_dir:
            output_path = Path(output_dir) / input_path.with_suffix('.docx').name
        else:
            output_path = None

        result = convert_md_to_docx(str(input_path), str(output_path) if output_path else None)
        results.append(result)
        print(f"Converted: {input_path} -> {result}")

    return results


def main():
    """Command line interface."""
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py <input.md> [output.docx]")
        print("       python md_to_docx.py <file1.md> <file2.md> ... --output-dir <dir>")
        sys.exit(1)

    args = sys.argv[1:]

    # Check for --output-dir flag for batch conversion
    if '--output-dir' in args:
        idx = args.index('--output-dir')
        output_dir = args[idx + 1]
        input_files = args[:idx]
        convert_multiple(input_files, output_dir)
    elif len(args) == 1:
        # Single file, output to same location
        result = convert_md_to_docx(args[0])
        print(f"Created: {result}")
    elif len(args) == 2:
        # Single file with explicit output
        result = convert_md_to_docx(args[0], args[1])
        print(f"Created: {result}")
    else:
        # Multiple files, output to same locations
        convert_multiple(args)


if __name__ == '__main__':
    main()
